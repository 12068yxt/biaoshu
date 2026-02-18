#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LangGraph 文档拆分 + 智能绘图 双Agent工作流

功能：读取docx按Heading 5拆分章节，LLM生成PPT兼容SVG，YAML配置管理提示词和风格

作者：AI Agent开发工程师
技术栈：LangGraph、python-docx、langchain-openai、PyYAML
"""

import os
import re
import json
import yaml
import time
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional, Tuple, Any, Callable
from string import Template
from datetime import datetime
from pathlib import Path

# 第三方库
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# LangGraph相关
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver

# LangChain相关
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage


# =============================================================================
# 数据模型定义
# =============================================================================

@dataclass
class Section:
    """
    文档章节数据类
    
    Attributes:
        index: 章节序号（0-based）
        title: Heading 5标题文本
        content: 合并后续所有正文段落（\n分隔）
        hierarchy_path: 上级标题路径，如"第一章>第二节>架构设计"
    """
    index: int
    title: str
    content: str
    hierarchy_path: str


@dataclass
class SVGResult:
    """
    SVG生成结果数据类
    
    Attributes:
        section_index: 对应章节索引
        section_title: 章节标题
        svg_content: SVG代码内容
        svg_path: SVG文件保存路径
        success: 是否生成成功
        error_message: 错误信息（失败时）
        timestamp: 生成时间戳
    """
    section_index: int
    section_title: str
    svg_content: str
    svg_path: str
    success: bool
    error_message: str = ""
    timestamp: str = field(default_factory=lambda: datetime.now().isoformat())


@dataclass
class WorkflowState:
    """
    LangGraph工作流状态类
    
    用于在各个环节之间传递数据
    """
    # 配置相关
    config_path: str = "prompts/standard.yaml"
    config_manager: Optional[Any] = None
    
    # 输入文件
    docx_path: str = ""
    
    # 拆分结果
    sections: List[Section] = field(default_factory=list)
    split_success: bool = False
    split_error: str = ""
    
    # 绘图进度
    current_section_idx: int = 0
    svg_results: List[SVGResult] = field(default_factory=list)
    
    # 最终结果
    report_path: str = "output/report.json"
    workflow_success: bool = False
    error_message: str = ""
    
    # 运行时数据
    output_dir: str = "output"


# =============================================================================
# 配置管理类
# =============================================================================

class ConfigManager:
    """
    配置管理类
    
    负责：
    1. YAML配置文件的加载和解析
    2. 提示词文件（system.txt, user.txt, examples.txt）的读取
    3. 提示词模板变量的渲染
    4. 配置热重载支持
    
    Attributes:
        config_path: YAML配置文件路径
        config_data: 解析后的配置数据
        file_mtimes: 文件修改时间记录（用于热重载）
    """
    
    def __init__(self, config_path: str = "prompts/standard.yaml"):
        """
        初始化配置管理器
        
        Args:
            config_path: YAML配置文件路径
        """
        self.config_path = config_path
        self.config_data: Dict[str, Any] = {}
        self.file_mtimes: Dict[str, float] = {}
        self._load_config()
    
    def _load_config(self) -> None:
        """
        加载YAML配置文件
        
        Raises:
            FileNotFoundError: 配置文件不存在
            yaml.YAMLError: YAML解析错误
        """
        if not os.path.exists(self.config_path):
            raise FileNotFoundError(f"配置文件不存在: {self.config_path}")
        
        with open(self.config_path, 'r', encoding='utf-8') as f:
            self.config_data = yaml.safe_load(f)
        
        # 记录文件修改时间
        self.file_mtimes[self.config_path] = os.path.getmtime(self.config_path)
        
        # 记录提示词文件修改时间
        prompts_config = self.config_data.get('prompts', {})
        for key in ['system_file', 'user_file', 'examples_file']:
            file_path = prompts_config.get(key)
            if file_path and os.path.exists(file_path):
                self.file_mtimes[file_path] = os.path.getmtime(file_path)
    
    def reload_if_changed(self) -> bool:
        """
        检查文件是否修改，如有修改则重新加载
        
        Returns:
            是否执行了重新加载
        """
        need_reload = False
        
        for file_path, last_mtime in self.file_mtimes.items():
            if os.path.exists(file_path):
                current_mtime = os.path.getmtime(file_path)
                if current_mtime > last_mtime:
                    need_reload = True
                    break
        
        if need_reload:
            self._load_config()
            return True
        return False
    
    def get_style_config(self) -> Dict[str, Any]:
        """
        获取样式配置
        
        Returns:
            样式配置字典，包含colors、bg_color、text_color、font等
        """
        return self.config_data.get('style', {})
    
    def get_llm_config(self) -> Dict[str, Any]:
        """
        获取LLM配置
        
        Returns:
            LLM配置字典，包含model、temperature、max_tokens等
        """
        return self.config_data.get('llm', {})
    
    def get_output_config(self) -> Dict[str, Any]:
        """
        获取输出配置
        
        Returns:
            输出配置字典，包含svg_dir、report_file等
        """
        return self.config_data.get('output', {})
    
    def load_prompts(self) -> Tuple[str, str]:
        """
        加载并组合提示词
        
        加载逻辑：
        1. 读取standard.yaml获取文件路径配置
        2. 分别读取system.txt、user.txt、examples.txt
        3. 拼接：system + "\n\n## 示例\n" + examples
        4. 返回(system, user)元组（未渲染变量）
        
        Returns:
            (system_prompt, user_prompt) 元组
        
        Raises:
            FileNotFoundError: 提示词文件不存在
        """
        prompts_config = self.config_data.get('prompts', {})
        
        # 获取文件路径
        system_file = prompts_config.get('system_file', 'prompts/system.txt')
        user_file = prompts_config.get('user_file', 'prompts/user.txt')
        examples_file = prompts_config.get('examples_file', 'prompts/examples.txt')
        
        # 读取system提示词
        if not os.path.exists(system_file):
            raise FileNotFoundError(f"系统提示词文件不存在: {system_file}")
        with open(system_file, 'r', encoding='utf-8') as f:
            system_prompt = f.read()
        
        # 读取user提示词模板
        if not os.path.exists(user_file):
            raise FileNotFoundError(f"用户提示词文件不存在: {user_file}")
        with open(user_file, 'r', encoding='utf-8') as f:
            user_prompt = f.read()
        
        # 读取examples并拼接到system
        if os.path.exists(examples_file):
            with open(examples_file, 'r', encoding='utf-8') as f:
                examples_content = f.read()
            if examples_content.strip():
                system_prompt = f"{system_prompt}\n\n## 示例\n{examples_content}"
        
        return system_prompt, user_prompt
    
    def render_prompts(
        self,
        title: str,
        content: str,
        hierarchy_path: str,
        **extra_vars
    ) -> Tuple[str, str]:
        """
        渲染提示词模板
        
        将变量填充到提示词模板中，生成最终的提示词
        
        Args:
            title: 章节标题
            content: 章节内容
            hierarchy_path: 层级路径
            **extra_vars: 额外变量
        
        Returns:
            (rendered_system, rendered_user) 元组
        """
        # 获取样式配置
        style = self.get_style_config()
        colors = style.get('colors', {})
        
        # 构建变量字典
        template_vars = {
            'title': title,
            'content': content,
            'hierarchy_path': hierarchy_path,
            'colors': colors,
            'bg_color': style.get('bg_color', '#FFFFFF'),
            'text_color': style.get('text_color', '#2C3E50'),
            'font': style.get('font', 'Arial, sans-serif'),
            **extra_vars
        }
        
        # 加载原始提示词
        system_template, user_template = self.load_prompts()
        
        # 渲染模板（使用string.Template）
        system_prompt = Template(system_template).safe_substitute(template_vars)
        user_prompt = Template(user_template).safe_substitute(template_vars)
        
        return system_prompt, user_prompt


# =============================================================================
# 文档拆分类
# =============================================================================

class DocumentSplitter:
    """
    文档拆分类
    
    负责：
    1. 解析docx文档
    2. 识别Heading 5标题
    3. 构建层级路径（Heading 1-4）
    4. 收集章节内容
    
    Attributes:
        docx_path: Word文档路径
    """
    
    def __init__(self, docx_path: str):
        """
        初始化文档拆分器
        
        Args:
            docx_path: Word文档路径
        """
        self.docx_path = docx_path
    
    def _get_paragraph_style(self, paragraph) -> Optional[str]:
        """
        获取段落的样式名称
        
        Args:
            paragraph: docx段落对象
        
        Returns:
            样式名称（如"Heading 1"），如果不是标题则返回None
        """
        style_name = paragraph.style.name if paragraph.style else None
        return style_name
    
    def _is_heading(self, style_name: Optional[str], level: int) -> bool:
        """
        检查样式是否为指定级别的标题
        
        Args:
            style_name: 样式名称
            level: 标题级别（1-5）
        
        Returns:
            是否匹配
        """
        if not style_name:
            return False
        # 支持多种命名格式："Heading 1", "Heading1", "标题 1"等
        patterns = [
            f"^Heading {level}$",
            f"^Heading{level}$",
            f"^标题 {level}$",
            f"^标题{level}$",
        ]
        return any(re.match(pattern, style_name, re.IGNORECASE) for pattern in patterns)
    
    def _is_any_heading(self, style_name: Optional[str]) -> bool:
        """
        检查样式是否为任意级别的标题（Heading 1-5）
        
        Args:
            style_name: 样式名称
        
        Returns:
            是否为标题样式
        """
        if not style_name:
            return False
        patterns = [
            r"^Heading [1-5]$",
            r"^Heading[1-5]$",
            r"^标题 [1-5]$",
            r"^标题[1-5]$",
        ]
        return any(re.match(pattern, style_name, re.IGNORECASE) for pattern in patterns)
    
    def split_by_heading5(self) -> List[Section]:
        """
        按Heading 5拆分文档
        
        拆分逻辑：
        1. 遍历文档所有段落
        2. 遇到"Heading 5"时，创建新章节，该标题作为章节标题
        3. 收集后续段落内容（非Heading样式），作为章节正文
        4. 遇到下一个"Heading 1-5"时，当前章节结束
        5. 层级路径：拼接当前章节所在的上级标题（Heading 1-4），用">"连接
        
        Returns:
            Section对象列表
        
        Raises:
            FileNotFoundError: 文档不存在
            Exception: 文档解析错误
        """
        if not os.path.exists(self.docx_path):
            raise FileNotFoundError(f"文档不存在: {self.docx_path}")
        
        # 加载文档
        doc = Document(self.docx_path)
        
        sections: List[Section] = []
        current_section: Optional[Dict] = None
        
        # 层级路径追踪
        heading_stack: Dict[int, str] = {}  # {level: title}
        
        for paragraph in doc.paragraphs:
            style_name = self._get_paragraph_style(paragraph)
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检查是否为Heading 1-5
            is_heading = False
            heading_level = 0
            
            for level in range(1, 6):
                if self._is_heading(style_name, level):
                    is_heading = True
                    heading_level = level
                    break
            
            if is_heading:
                # 更新层级路径
                heading_stack[heading_level] = text
                # 清除更低级别的标题
                for l in list(heading_stack.keys()):
                    if l > heading_level:
                        del heading_stack[l]
                
                # 如果是Heading 5，创建新章节
                if heading_level == 5:
                    # 先保存当前章节（如果存在）
                    if current_section:
                        sections.append(Section(
                            index=len(sections),
                            title=current_section['title'],
                            content='\n'.join(current_section['content']),
                            hierarchy_path=current_section['hierarchy_path']
                        ))
                    
                    # 构建层级路径（Heading 1-4）
                    hierarchy_parts = []
                    for l in range(1, 5):
                        if l in heading_stack:
                            hierarchy_parts.append(heading_stack[l])
                    hierarchy_path = '>'.join(hierarchy_parts) if hierarchy_parts else ""
                    
                    # 创建新章节
                    current_section = {
                        'title': text,
                        'content': [],
                        'hierarchy_path': hierarchy_path
                    }
            else:
                # 非标题段落，添加到当前章节内容
                if current_section is not None:
                    current_section['content'].append(text)
        
        # 保存最后一个章节
        if current_section:
            sections.append(Section(
                index=len(sections),
                title=current_section['title'],
                content='\n'.join(current_section['content']),
                hierarchy_path=current_section['hierarchy_path']
            ))
        
        return sections


# =============================================================================
# 智能绘图类
# =============================================================================

class SmartDrawer:
    """
    智能绘图类
    
    负责：
    1. 提示词组装（调用ConfigManager）
    2. LLM调用（带重试机制）
    3. SVG验证和修复
    4. 失败回退（生成极简备用SVG）
    
    Attributes:
        config_manager: 配置管理器实例
        llm: LangChain LLM实例
        retry_times: 重试次数
    """
    
    def __init__(self, config_manager: ConfigManager):
        """
        初始化智能绘图器
        
        Args:
            config_manager: 配置管理器实例
        """
        self.config_manager = config_manager
        self.llm_config = config_manager.get_llm_config()
        self.retry_times = self.llm_config.get('retry_times', 2)
        
        # 初始化LLM
        self.llm = ChatOpenAI(
            model=self.llm_config.get('model', 'gpt-4o'),
            temperature=self.llm_config.get('temperature', 0.3),
            max_tokens=self.llm_config.get('max_tokens', 4000)
        )
    
    def _call_llm(self, system_prompt: str, user_prompt: str) -> str:
        """
        调用LLM生成SVG
        
        Args:
            system_prompt: 系统提示词
            user_prompt: 用户提示词
        
        Returns:
            LLM生成的SVG代码
        
        Raises:
            Exception: LLM调用失败
        """
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_prompt)
        ]
        
        response = self.llm.invoke(messages)
        return response.content
    
    def _extract_svg(self, content: str) -> Optional[str]:
        """
        从LLM响应中提取SVG代码
        
        Args:
            content: LLM原始响应
        
        Returns:
            提取的SVG代码，如果没有找到则返回None
        """
        # 查找<svg>标签
        svg_match = re.search(r'<svg[^>]*>.*?</svg>', content, re.DOTALL | re.IGNORECASE)
        if svg_match:
            svg_content = svg_match.group(0)
            
            # 检查是否有XML声明
            if not svg_content.startswith('<?xml'):
                svg_content = '<?xml version="1.0" encoding="UTF-8"?>\n' + svg_content
            
            return svg_content
        
        return None
    
    def _validate_svg(self, svg_content: str) -> Tuple[bool, str]:
        """
        验证SVG代码的完整性和正确性
        
        验证项：
        1. 包含<svg>标签
        2. 包含xmlns命名空间
        3. 包含viewBox属性
        4. 标签正确闭合
        
        Args:
            svg_content: SVG代码
        
        Returns:
            (是否有效, 错误信息) 元组
        """
        # 检查<svg>标签
        if '<svg' not in svg_content.lower():
            return False, "缺少<svg>标签"
        
        # 检查xmlns
        if 'xmlns="http://www.w3.org/2000/svg"' not in svg_content:
            return False, "缺少xmlns命名空间"
        
        # 检查viewBox
        if 'viewBox' not in svg_content:
            return False, "缺少viewBox属性"
        
        # 检查标签闭合
        svg_open_count = len(re.findall(r'<svg[^>]*>', svg_content, re.IGNORECASE))
        svg_close_count = len(re.findall(r'</svg>', svg_content, re.IGNORECASE))
        if svg_open_count != svg_close_count:
            return False, "<svg>标签未正确闭合"
        
        return True, ""
    
    def _fix_svg(self, svg_content: str) -> str:
        """
        自动修复SVG代码中的常见问题
        
        修复项：
        1. 添加缺失的xmlns
        2. 添加缺失的viewBox
        3. 确保XML声明
        
        Args:
            svg_content: 原始SVG代码
        
        Returns:
            修复后的SVG代码
        """
        # 确保XML声明
        if not svg_content.startswith('<?xml'):
            svg_content = '<?xml version="1.0" encoding="UTF-8"?>\n' + svg_content
        
        # 添加xmlns（如果不存在）
        if 'xmlns="http://www.w3.org/2000/svg"' not in svg_content:
            svg_content = svg_content.replace('<svg', '<svg xmlns="http://www.w3.org/2000/svg"', 1)
        
        # 添加viewBox（如果不存在）
        if 'viewBox' not in svg_content:
            # 尝试提取width和height
            width_match = re.search(r'width=["\'](\d+)["\']', svg_content)
            height_match = re.search(r'height=["\'](\d+)["\']', svg_content)
            
            if width_match and height_match:
                width = width_match.group(1)
                height = height_match.group(1)
                svg_content = svg_content.replace('<svg', f'<svg viewBox="0 0 {width} {height}"', 1)
            else:
                # 默认viewBox
                svg_content = svg_content.replace('<svg', '<svg viewBox="0 0 800 600"', 1)
        
        return svg_content
    
    def _generate_fallback_svg(self, title: str) -> str:
        """
        生成极简备用SVG
        
        当LLM调用失败时，生成一个简单的占位SVG
        
        Args:
            title: 章节标题
        
        Returns:
            备用SVG代码
        """
        style = self.config_manager.get_style_config()
        bg_color = style.get('bg_color', '#FFFFFF')
        text_color = style.get('text_color', '#2C3E50')
        primary_color = style.get('colors', {}).get('primary', '#4A90D9')
        
        # 截断标题（避免过长）
        display_title = title[:50] + "..." if len(title) > 50 else title
        
        fallback_svg = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 600" width="800" height="600">
  <rect width="800" height="600" fill="{bg_color}"/>
  <rect x="50" y="50" width="700" height="500" rx="10" fill="none" stroke="{primary_color}" stroke-width="2"/>
  <text x="400" y="280" text-anchor="middle" font-family="Arial, sans-serif" font-size="24" fill="{text_color}">
    {display_title}
  </text>
  <text x="400" y="320" text-anchor="middle" font-family="Arial, sans-serif" font-size="14" fill="{primary_color}">
    [SVG生成失败 - 备用图表]
  </text>
</svg>'''
        return fallback_svg
    
    def draw(
        self,
        section: Section,
        output_dir: str = "output/svgs"
    ) -> SVGResult:
        """
        为章节生成SVG图表
        
        流程：
        1. 组装提示词（调用ConfigManager.render_prompts）
        2. 调用LLM（带重试机制）
        3. 提取和验证SVG
        4. 保存SVG文件
        5. 失败时生成备用SVG
        
        Args:
            section: 章节对象
            output_dir: SVG输出目录
        
        Returns:
            SVGResult对象
        """
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 准备输出路径
        safe_title = re.sub(r'[^\w\u4e00-\u9fff]', '_', section.title)[:30]
        svg_filename = f"section_{section.index:03d}_{safe_title}.svg"
        svg_path = os.path.join(output_dir, svg_filename)
        
        # 渲染提示词
        system_prompt, user_prompt = self.config_manager.render_prompts(
            title=section.title,
            content=section.content,
            hierarchy_path=section.hierarchy_path
        )
        
        # 重试机制
        last_error = ""
        for attempt in range(self.retry_times + 1):
            try:
                # 调用LLM
                llm_response = self._call_llm(system_prompt, user_prompt)
                
                # 提取SVG
                svg_content = self._extract_svg(llm_response)
                
                if svg_content is None:
                    last_error = "无法从LLM响应中提取SVG代码"
                    continue
                
                # 验证SVG
                is_valid, error_msg = self._validate_svg(svg_content)
                
                if not is_valid:
                    # 尝试修复
                    svg_content = self._fix_svg(svg_content)
                    is_valid, error_msg = self._validate_svg(svg_content)
                    
                    if not is_valid:
                        last_error = f"SVG验证失败: {error_msg}"
                        continue
                
                # 保存SVG文件
                with open(svg_path, 'w', encoding='utf-8') as f:
                    f.write(svg_content)
                
                return SVGResult(
                    section_index=section.index,
                    section_title=section.title,
                    svg_content=svg_content,
                    svg_path=svg_path,
                    success=True
                )
                
            except Exception as e:
                last_error = str(e)
                if attempt < self.retry_times:
                    time.sleep(1)  # 短暂延迟后重试
                    continue
        
        # 所有重试失败，生成备用SVG
        fallback_svg = self._generate_fallback_svg(section.title)
        
        with open(svg_path, 'w', encoding='utf-8') as f:
            f.write(fallback_svg)
        
        return SVGResult(
            section_index=section.index,
            section_title=section.title,
            svg_content=fallback_svg,
            svg_path=svg_path,
            success=False,
            error_message=f"LLM调用失败（重试{self.retry_times}次）: {last_error}"
        )


# =============================================================================
# LangGraph工作流节点函数
# =============================================================================

def initialize(state: WorkflowState) -> WorkflowState:
    """
    初始化节点
    
    职责：
    1. 加载配置，创建ConfigManager
    2. 创建输出目录
    3. 初始化状态
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    print("[初始化] 开始加载配置...")
    
    try:
        # 创建配置管理器
        state.config_manager = ConfigManager(state.config_path)
        
        # 获取输出配置
        output_config = state.config_manager.get_output_config()
        state.output_dir = os.path.dirname(output_config.get('report_file', 'output/report.json'))
        svg_dir = output_config.get('svg_dir', 'output/svgs')
        
        # 创建输出目录
        os.makedirs(state.output_dir, exist_ok=True)
        os.makedirs(svg_dir, exist_ok=True)
        
        state.report_path = output_config.get('report_file', 'output/report.json')
        
        print(f"[初始化] 配置加载成功")
        print(f"[初始化] 输出目录: {state.output_dir}")
        print(f"[初始化] SVG目录: {svg_dir}")
        
    except Exception as e:
        state.workflow_success = False
        state.error_message = f"初始化失败: {str(e)}"
        print(f"[初始化] 错误: {e}")
    
    return state


def split_document(state: WorkflowState) -> WorkflowState:
    """
    文档拆分节点
    
    职责：
    1. 使用DocumentSplitter解析docx
    2. 按Heading 5拆分章节
    3. 输出List[Section]
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    print(f"[文档拆分] 开始解析文档: {state.docx_path}")
    
    try:
        # 创建文档拆分器
        splitter = DocumentSplitter(state.docx_path)
        
        # 拆分文档
        sections = splitter.split_by_heading5()
        
        state.sections = sections
        state.split_success = len(sections) > 0
        
        if state.split_success:
            print(f"[文档拆分] 成功拆分 {len(sections)} 个章节")
            for section in sections:
                print(f"  [{section.index}] {section.title}")
                print(f"      路径: {section.hierarchy_path}")
                print(f"      内容长度: {len(section.content)} 字符")
        else:
            state.split_error = "未找到Heading 5章节"
            print("[文档拆分] 警告: 未找到任何Heading 5章节")
            
    except Exception as e:
        state.split_success = False
        state.split_error = str(e)
        print(f"[文档拆分] 错误: {e}")
    
    return state


def check_split_result(state: WorkflowState) -> str:
    """
    拆分结果检查条件
    
    条件分支：
    - 成功→"prepare_draw"
    - 失败/空→"handle_error"
    
    Args:
        state: 工作流状态
    
    Returns:
        下一个节点名称
    """
    if state.split_success and len(state.sections) > 0:
        return "prepare_draw"
    else:
        return "handle_error"


def prepare_draw(state: WorkflowState) -> WorkflowState:
    """
    准备绘图节点
    
    职责：
    1. 设置current_section_idx
    2. 打印进度信息
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    total = len(state.sections)
    current = state.current_section_idx
    
    if current < total:
        section = state.sections[current]
        print(f"\n[准备绘图] 进度: {current + 1}/{total}")
        print(f"[准备绘图] 当前章节: {section.title}")
        print(f"[准备绘图] 层级路径: {section.hierarchy_path}")
    
    return state


def draw_svg(state: WorkflowState) -> WorkflowState:
    """
    SVG绘图节点
    
    职责：
    1. 使用SmartDrawer获取提示词
    2. 调用LLM生成SVG
    3. 验证保存SVG
    4. 记录结果
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    if state.current_section_idx >= len(state.sections):
        return state
    
    section = state.sections[state.current_section_idx]
    print(f"[智能绘图] 正在生成: {section.title}")
    
    try:
        # 创建智能绘图器
        drawer = SmartDrawer(state.config_manager)
        
        # 获取SVG输出目录
        output_config = state.config_manager.get_output_config()
        svg_dir = output_config.get('svg_dir', 'output/svgs')
        
        # 生成SVG
        result = drawer.draw(section, svg_dir)
        
        # 记录结果
        state.svg_results.append(result)
        
        if result.success:
            print(f"[智能绘图] 成功: {result.svg_path}")
        else:
            print(f"[智能绘图] 失败（使用备用）: {result.error_message}")
            
    except Exception as e:
        # 记录失败结果
        error_result = SVGResult(
            section_index=section.index,
            section_title=section.title,
            svg_content="",
            svg_path="",
            success=False,
            error_message=str(e)
        )
        state.svg_results.append(error_result)
        print(f"[智能绘图] 错误: {e}")
    
    return state


def check_continue(state: WorkflowState) -> str:
    """
    继续检查条件
    
    条件分支：
    - 还有章节→"prepare_draw"（循环）
    - 全部完成→"generate_report"
    
    Args:
        state: 工作流状态
    
    Returns:
        下一个节点名称
    """
    # 递增索引
    state.current_section_idx += 1
    
    if state.current_section_idx < len(state.sections):
        return "prepare_draw"
    else:
        return "generate_report"


def generate_report(state: WorkflowState) -> WorkflowState:
    """
    生成报告节点
    
    职责：
    1. 汇总List[SVGResult]
    2. 写JSON报告
    3. 打印统计信息
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    print("\n[生成报告] 开始汇总结果...")
    
    try:
        # 统计信息
        total = len(state.svg_results)
        success_count = sum(1 for r in state.svg_results if r.success)
        failed_count = total - success_count
        
        # 构建报告数据
        report = {
            "timestamp": datetime.now().isoformat(),
            "source_document": state.docx_path,
            "statistics": {
                "total_sections": total,
                "success_count": success_count,
                "failed_count": failed_count,
                "success_rate": f"{success_count/total*100:.1f}%" if total > 0 else "0%"
            },
            "sections": []
        }
        
        # 添加章节详情
        for result in state.svg_results:
            report["sections"].append({
                "index": result.section_index,
                "title": result.section_title,
                "svg_path": result.svg_path,
                "success": result.success,
                "error_message": result.error_message if not result.success else "",
                "timestamp": result.timestamp
            })
        
        # 保存报告
        with open(state.report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        state.workflow_success = True
        
        print(f"[生成报告] 报告已保存: {state.report_path}")
        print(f"[生成报告] 统计: 总计{total} | 成功{success_count} | 失败{failed_count}")
        
    except Exception as e:
        state.workflow_success = False
        state.error_message = f"生成报告失败: {str(e)}"
        print(f"[生成报告] 错误: {e}")
    
    return state


def handle_error(state: WorkflowState) -> WorkflowState:
    """
    错误处理节点
    
    职责：
    1. 记录错误信息
    2. 优雅退出
    
    Args:
        state: 工作流状态
    
    Returns:
        更新后的状态
    """
    print(f"\n[错误处理] 工作流遇到错误")
    
    if state.split_error:
        state.error_message = f"文档拆分错误: {state.split_error}"
    elif not state.error_message:
        state.error_message = "未知错误"
    
    print(f"[错误处理] 错误信息: {state.error_message}")
    print(f"[错误处理] 工作流已终止")
    
    state.workflow_success = False
    
    return state


# =============================================================================
# 构建LangGraph工作流
# =============================================================================

def build_workflow() -> StateGraph:
    """
    构建LangGraph工作流
    
    工作流图结构：
    
                    ┌─────────────┐
                    │ initialize  │
                    └──────┬──────┘
                           │
                           ▼
                    ┌─────────────┐
                    │split_document│
                    └──────┬──────┘
                           │
              ┌────────────┴────────────┐
              │                         │
              ▼                         ▼
    ┌─────────────────┐      ┌─────────────────┐
    │  check_split    │      │   handle_error  │
    │  (条件检查)      │      │   (错误处理)     │
    └────────┬────────┘      └─────────────────┘
             │ 成功
             ▼
    ┌─────────────────┐
    │  prepare_draw   │◄────────────────┐
    │  (准备绘图)      │                 │
    └────────┬────────┘                 │
             │                          │
             ▼                          │
    ┌─────────────────┐                 │
    │    draw_svg     │                 │
    │  (智能绘图)      │                 │
    └────────┬────────┘                 │
             │                          │
             ▼                          │
    ┌─────────────────┐                 │
    │  check_continue │─────────────────┘
    │  (继续检查)      │（还有章节）
    └────────┬────────┘
             │ 完成
             ▼
    ┌─────────────────┐
    │ generate_report │
    │   (生成报告)     │
    └─────────────────┘
    
    Returns:
        编译后的StateGraph
    """
    # 创建工作流图
    workflow = StateGraph(WorkflowState)
    
    # 添加节点
    workflow.add_node("initialize", initialize)
    workflow.add_node("split_document", split_document)
    workflow.add_node("prepare_draw", prepare_draw)
    workflow.add_node("draw_svg", draw_svg)
    workflow.add_node("generate_report", generate_report)
    workflow.add_node("handle_error", handle_error)
    
    # 设置入口节点
    workflow.set_entry_point("initialize")
    
    # 添加边
    workflow.add_edge("initialize", "split_document")
    
    # 条件分支：拆分结果检查
    workflow.add_conditional_edges(
        "split_document",
        check_split_result,
        {
            "prepare_draw": "prepare_draw",
            "handle_error": "handle_error"
        }
    )
    
    # 准备绘图 -> 智能绘图
    workflow.add_edge("prepare_draw", "draw_svg")
    
    # 条件分支：继续检查（循环或完成）
    workflow.add_conditional_edges(
        "draw_svg",
        check_continue,
        {
            "prepare_draw": "prepare_draw",
            "generate_report": "generate_report"
        }
    )
    
    # 结束节点
    workflow.add_edge("generate_report", END)
    workflow.add_edge("handle_error", END)
    
    # 编译工作流（使用内存检查点）
    memory = MemorySaver()
    compiled_workflow = workflow.compile(checkpointer=memory)
    
    return compiled_workflow


# =============================================================================
# 主运行入口
# =============================================================================

def run_workflow(
    docx_path: str,
    config_path: str = "prompts/standard.yaml",
    verbose: bool = True
) -> WorkflowState:
    """
    运行完整工作流
    
    Args:
        docx_path: Word文档路径
        config_path: 配置文件路径
        verbose: 是否打印详细日志
    
    Returns:
        最终工作流状态
    
    Example:
        >>> result = run_workflow("document.docx")
        >>> print(f"成功: {result.workflow_success}")
        >>> print(f"生成SVG: {len(result.svg_results)} 个")
    """
    # 检查文档是否存在
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文档不存在: {docx_path}")
    
    # 创建工作流
    workflow = build_workflow()
    
    # 初始化状态
    initial_state = WorkflowState(
        docx_path=docx_path,
        config_path=config_path
    )
    
    # 运行工作流
    print("=" * 60)
    print("LangGraph 文档拆分 + 智能绘图 工作流")
    print("=" * 60)
    print(f"文档: {docx_path}")
    print(f"配置: {config_path}")
    print("=" * 60)
    
    # 执行工作流
    final_state = workflow.invoke(
        initial_state,
        config={"configurable": {"thread_id": "svg_workflow_" + str(int(time.time()))}}
    )
    
    print("\n" + "=" * 60)
    print("工作流执行完成")
    print("=" * 60)
    print(f"状态: {'成功' if final_state.workflow_success else '失败'}")
    if final_state.error_message:
        print(f"错误: {final_state.error_message}")
    print(f"生成SVG: {len(final_state.svg_results)} 个")
    if final_state.report_path:
        print(f"报告: {final_state.report_path}")
    print("=" * 60)
    
    return final_state


# =============================================================================
# 辅助函数
# =============================================================================

def create_sample_docx(output_path: str = "examples/sample.docx") -> str:
    """
    创建示例Word文档用于测试
    
    创建一个包含多级标题和内容的文档，用于测试拆分功能
    
    Args:
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    doc.add_heading('系统架构设计文档', 0)
    
    # 第一章
    doc.add_heading('第一章 系统概述', level=1)
    doc.add_paragraph('本章介绍系统的整体架构和设计原则。')
    
    doc.add_heading('1.1 设计目标', level=2)
    doc.add_paragraph('系统采用微服务架构，实现高可用、可扩展的设计目标。')
    
    doc.add_heading('1.1.1 高可用性', level=3)
    doc.add_paragraph('通过多副本部署和负载均衡实现99.99%的可用性。')
    
    doc.add_heading('1.1.1.1 模块划分原则', level=5)
    doc.add_paragraph('系统按照业务领域进行模块划分，每个模块独立部署。')
    doc.add_paragraph('核心模块包括：用户服务、订单服务、支付服务。')
    doc.add_paragraph('模块间通过消息队列进行异步通信。')
    
    doc.add_heading('1.1.1.2 依赖关系设计', level=5)
    doc.add_paragraph('模块依赖遵循单向依赖原则，避免循环依赖。')
    doc.add_paragraph('基础服务不依赖业务服务，业务服务可依赖基础服务。')
    
    doc.add_heading('1.2 技术选型', level=2)
    
    doc.add_heading('1.2.1.1 后端技术栈', level=5)
    doc.add_paragraph('后端采用Spring Boot框架，数据库使用MySQL和Redis。')
    doc.add_paragraph('缓存层使用Redis集群，消息队列使用RabbitMQ。')
    
    doc.add_heading('1.2.1.2 前端技术栈', level=5)
    doc.add_paragraph('前端采用Vue3 + TypeScript技术栈。')
    doc.add_paragraph('UI组件库使用Element Plus，状态管理使用Pinia。')
    
    # 第二章
    doc.add_heading('第二章 数据架构', level=1)
    
    doc.add_heading('2.1.1.1 数据库设计', level=5)
    doc.add_paragraph('采用分库分表策略，按用户ID进行数据分片。')
    doc.add_paragraph('主从复制实现读写分离，提高查询性能。')
    
    doc.add_heading('2.1.1.2 缓存策略', level=5)
    doc.add_paragraph('热点数据缓存，设置合理的过期时间。')
    doc.add_paragraph('使用缓存穿透、击穿、雪崩的防护策略。')
    
    # 保存文档
    doc.save(output_path)
    
    print(f"示例文档已创建: {output_path}")
    return output_path


def main():
    """
    主函数 - 命令行入口
    
    用法:
        python svg_workflow.py <docx文件路径>
        python svg_workflow.py --sample  # 创建并运行示例文档
    """
    import sys
    
    # 切换到脚本所在目录（确保相对路径正确）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)
    
    if len(sys.argv) > 1:
        if sys.argv[1] == '--sample':
            # 创建并运行示例
            docx_path = create_sample_docx()
            run_workflow(docx_path)
        else:
            # 运行指定文档
            run_workflow(sys.argv[1])
    else:
        # 默认运行示例
        print("用法:")
        print("  python svg_workflow.py <docx文件路径>")
        print("  python svg_workflow.py --sample  # 创建并运行示例文档")
        print("\n正在创建并运行示例文档...")
        docx_path = create_sample_docx()
        run_workflow(docx_path)


if __name__ == "__main__":
    main()