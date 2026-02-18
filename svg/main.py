#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SVG 工作流主入口

LangGraph 文档拆分 + 智能绘图工作流

用法:
    python main.py <docx文件路径>
    python main.py --sample  # 创建并运行示例文档
"""

import os
import sys
from datetime import datetime

# 添加 src 到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from src.agents import build_workflow, create_initial_state


def run_workflow(
    docx_path: str,
    config_path: str = "src/config/standard.yaml",
) -> dict:
    """
    运行完整工作流

    Args:
        docx_path: Word 文档路径
        config_path: 配置文件路径

    Returns:
        最终工作流状态字典
    """
    # 检查文档是否存在
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文档不存在: {docx_path}")

    # 创建工作流
    workflow = build_workflow()

    # 初始化状态
    initial_state = create_initial_state(
        docx_path=docx_path,
        config_path=config_path,
    )

    # 运行工作流
    print("=" * 60)
    print("LangGraph 文档拆分 + 智能绘图 工作流")
    print("=" * 60)
    print(f"文档: {docx_path}")
    print(f"配置: {config_path}")
    print(f"线程: {initial_state['thread_id']}")
    print("=" * 60)

    # 执行工作流
    final_state = workflow.invoke(
        initial_state,
        config={"configurable": {"thread_id": initial_state["thread_id"]}}
    )

    print("\n" + "=" * 60)
    print("工作流执行完成")
    print("=" * 60)
    print(f"状态: {'成功' if final_state['workflow_success'] else '失败'}")
    if final_state['error_message']:
        print(f"错误: {final_state['error_message']}")
    print(f"生成SVG: {len(final_state['svg_results'])} 个")
    if final_state['report_path']:
        print(f"报告: {final_state['report_path']}")
    print("=" * 60)

    return final_state


def create_sample_docx(output_path: str = "examples/sample.docx") -> str:
    """
    创建示例 Word 文档用于测试

    Args:
        output_path: 输出文件路径

    Returns:
        输出文件路径
    """
    from docx import Document

    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 创建文档
    doc = Document()

    # 添加标题
    doc.add_heading('系统架构设计文档', 0)

    # 第一章
    doc.add_heading('第一章 系统概述', level=1)
    doc.add_paragraph('本章介绍系统的整体架构和设计原则。')

    doc.add_heading('1.1 设计目标', level=2)
    doc.add_paragraph('系统采用微服务架构，实现高可用、可扩展的设计目标。')

    doc.add_heading('1.1.1 高可用性', level=3)
    doc.add_paragraph('通过多副本部署和负载均衡实现99.99%的可用性。')

    doc.add_heading('1.1.1.1 模块划分原则', level=5)
    doc.add_paragraph('系统按照业务领域进行模块划分，每个模块独立部署。')
    doc.add_paragraph('核心模块包括：用户服务、订单服务、支付服务。')
    doc.add_paragraph('模块间通过消息队列进行异步通信。')

    doc.add_heading('1.1.1.2 依赖关系设计', level=5)
    doc.add_paragraph('模块依赖遵循单向依赖原则，避免循环依赖。')
    doc.add_paragraph('基础服务不依赖业务服务，业务服务可依赖基础服务。')

    doc.add_heading('1.2 技术选型', level=2)

    doc.add_heading('1.2.1.1 后端技术栈', level=5)
    doc.add_paragraph('后端采用Spring Boot框架，数据库使用MySQL和Redis。')
    doc.add_paragraph('缓存层使用Redis集群，消息队列使用RabbitMQ。')

    doc.add_heading('1.2.1.2 前端技术栈', level=5)
    doc.add_paragraph('前端采用Vue3 + TypeScript技术栈。')
    doc.add_paragraph('UI组件库使用Element Plus，状态管理使用Pinia。')

    # 第二章
    doc.add_heading('第二章 数据架构', level=1)

    doc.add_heading('2.1.1.1 数据库设计', level=5)
    doc.add_paragraph('采用分库分表策略，按用户ID进行数据分片。')
    doc.add_paragraph('主从复制实现读写分离，提高查询性能。')

    doc.add_heading('2.1.1.2 缓存策略', level=5)
    doc.add_paragraph('热点数据缓存，设置合理的过期时间。')
    doc.add_paragraph('使用缓存穿透、击穿、雪崩的防护策略。')

    # 保存文档
    doc.save(output_path)

    print(f"示例文档已创建: {output_path}")
    return output_path


def main():
    """主函数 - 命令行入口"""
    # 切换到脚本所在目录（确保相对路径正确）
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)

    if len(sys.argv) > 1:
        if sys.argv[1] == '--sample':
            # 创建并运行示例
            docx_path = create_sample_docx()
            run_workflow(docx_path)
        else:
            # 运行指定文档
            run_workflow(sys.argv[1])
    else:
        # 默认运行示例
        print("用法:")
        print("  python main.py <docx文件路径>")
        print("  python main.py --sample  # 创建并运行示例文档")
        print("\n正在创建并运行示例文档...")
        docx_path = create_sample_docx()
        run_workflow(docx_path)


if __name__ == "__main__":
    main()
